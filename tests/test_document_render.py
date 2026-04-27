from __future__ import annotations

import zipfile
from datetime import date
from io import BytesIO
from pathlib import Path

import pytest
from docxtpl import DocxTemplate
from docx import Document

from app.config import AppConfig
from app.models import (
    DocumentFieldType,
    DocumentTemplate,
    DocumentTemplateField,
    Project,
)
from app.services.storage import (
    StorageService,
    build_render_context,
    inspect_docx_tags,
)


def _make_storage(tmp_path: Path) -> StorageService:
    config = AppConfig.from_root(tmp_path)
    return StorageService(config)


def _make_docx(content_paragraphs: list[str]) -> bytes:
    document = Document()
    for line in content_paragraphs:
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_inspect_docx_tags_finds_jinja_variables(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(
        _make_docx(["Part: {{ part_number }}", "Supplier: {{ supplier_name }}", "Always-available: {{ project_name }}"])
    )
    tags = inspect_docx_tags(docx_path)
    assert set(tags) == {"part_number", "supplier_name", "project_name"}


def test_inspect_docx_tags_handles_loops_and_dotted_paths(tmp_path: Path) -> None:
    docx_path = tmp_path / "loop.docx"
    docx_path.write_bytes(
        _make_docx(["{% for m in milestones %}{{ m.title }} {% endfor %}", "Just text"])
    )
    tags = inspect_docx_tags(docx_path)
    assert "milestones" in tags
    # m.title should not appear as separate top-level tag — only "m"
    assert "title" not in tags


def test_render_document_template_substitutes_field_values(tmp_path: Path) -> None:
    storage = _make_storage(tmp_path)

    project = Project(slug="acme-rfq", name="ACME RFQ", description="Demo")
    storage.save_project(project, sections={"content": "", "change_requests": "", "roadblocks": "", "notes": ""})

    template = DocumentTemplate(
        slug="quote",
        name="Quote",
        fields=[
            DocumentTemplateField(key="part_number", label="Part No.", value="ABC-12345"),
            DocumentTemplateField(key="supplier_name", label="Supplier", value="Acme Corp"),
        ],
    )
    storage.save_document_template(template)

    docx_bytes = _make_docx(
        ["Part: {{ part_number }}", "Supplier: {{ supplier_name }}", "Project: {{ project_name }}"]
    )
    storage.save_document_template_file("quote", "quote.docx", docx_bytes)

    data, filename = storage.render_document_template("quote", "acme-rfq")
    assert filename.endswith(".docx")
    assert filename.startswith("quote__acme-rfq")

    # parse the rendered docx and confirm substitutions
    doc = Document(BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Part: ABC-12345" in text
    assert "Supplier: Acme Corp" in text
    assert "Project: ACME RFQ" in text


def test_render_without_uploaded_file_raises(tmp_path: Path) -> None:
    storage = _make_storage(tmp_path)
    project = Project(slug="empty-proj", name="Empty")
    storage.save_project(project, sections={"content": "", "change_requests": "", "roadblocks": "", "notes": ""})
    template = DocumentTemplate(slug="orphan", name="Orphan", fields=[])
    storage.save_document_template(template)

    with pytest.raises(FileNotFoundError):
        storage.render_document_template("orphan", "empty-proj")


def test_build_render_context_exposes_project_metadata(tmp_path: Path) -> None:
    project = Project(
        slug="ctx",
        name="Context Demo",
        description="What we test",
        start_date=date(2026, 1, 1),
        end_date=date(2026, 6, 30),
    )
    template = DocumentTemplate(slug="t", name="t", fields=[
        DocumentTemplateField(key="custom", label="Custom", value="v"),
    ])
    ctx = build_render_context(template, project)
    assert ctx["project_name"] == "Context Demo"
    assert ctx["project_start_date"] == "2026-01-01"
    assert ctx["custom"] == "v"
    assert ctx["today"]
